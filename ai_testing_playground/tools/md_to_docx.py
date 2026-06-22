"""Minimal Markdown -> Word (.docx) converter for the docs in this repo.

Handles the constructs used by docs/*.md: ATX headings, paragraphs, blockquote callouts,
fenced code blocks, pipe tables, ordered/unordered lists (incl. task items), horizontal
rules, and inline **bold** / *italic* / `code` / [links](url).

Usage:
    python tools/md_to_docx.py <input.md> <output.docx> [--title "Document Title"]

Requires: python-docx  (pip install python-docx)
"""
import argparse
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


CODE_FONT = "Consolas"
CODE_SHADE = "F2F3F5"
QUOTE_SHADE = "EAE6FF"
INLINE_RE = re.compile(
    r"(`[^`]+`)"                 # inline code
    r"|(\*\*[^*]+\*\*)"          # bold
    r"|(\[[^\]]+\]\([^)]+\))"    # link
    r"|(\*[^*]+\*)"              # italic
)


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_inline(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url)
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B5FFF")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_code_block(doc, lines):
    for line in lines or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _shade(p, CODE_SHADE)
        run = p.add_run(line if line else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(9)


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, cell_text in enumerate(row):
            para = cells[c_idx].paragraphs[0]
            add_inline(para, cell_text.strip())
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
    return table


def parse_table_block(block):
    rows = []
    for line in block:
        if re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def convert(md_path, docx_path, title=None):
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code)
            continue

        # Table (a pipe line followed by a separator line)
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$", lines[i + 1]):
            block = []
            while i < n and "|" in lines[i] and lines[i].strip():
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table_block(block))
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "B0B0B0")
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading("", level=min(level, 4))
            add_inline(heading, m.group(2).strip())
            i += 1
            continue

        # Blockquote / callout (consecutive > lines)
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            text = " ".join(q for q in quote_lines if q.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(10)
            _shade(p, QUOTE_SHADE)
            add_inline(p, text)
            continue

        # Task list item
        m = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)$", stripped)
        if m:
            checked = m.group(1).lower() == "x"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(("☑ " if checked else "☐ "))
            add_inline(p, m.group(2))
            i += 1
            continue

        # Unordered list item
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            i += 1
            continue

        # Ordered list item
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--title", default=None)
    args = ap.parse_args()
    convert(args.input, args.output, args.title)
